from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

doc = Document()

# ── 기본 스타일 설정 ──────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = '맑은 고딕'
style.font.size = Pt(10)
style._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

def set_heading_style(paragraph, level):
    colors = {1: '1F3864', 2: '2E74B5', 3: '2E74B5'}
    sizes  = {1: 18, 2: 14, 3: 12}
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run.font.size = Pt(sizes.get(level, 11))
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(colors.get(level, '000000'))
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

def add_table_from_md(doc, lines, start):
    """마크다운 테이블을 docx 테이블로 변환"""
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith('|'):
        row = [c.strip() for c in lines[i].strip().strip('|').split('|')]
        rows.append(row)
        i += 1
    # 구분선(---|---) 제거
    rows = [r for r in rows if not all(re.match(r'^-+$', c.replace(':', '').replace('-','').strip() or '-') for c in r)]

    if not rows:
        return i

    col_count = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = 'Table Grid'

    for ri, row in enumerate(rows):
        for ci, cell_text in enumerate(row):
            if ci >= col_count:
                break
            cell = table.cell(ri, ci)
            cell.text = cell_text
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.runs[0] if p.runs else p.add_run(cell_text)
            run.font.name = '맑은 고딕'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
            run.font.size = Pt(9)
            if ri == 0:
                run.font.bold = True
                # 헤더 배경색
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'D6E4F0')
                tcPr.append(shd)
    doc.add_paragraph()
    return i

def add_code_block(doc, code_lines):
    """코드 블록 스타일로 추가"""
    code_text = '\n'.join(code_lines)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x1F, 0x3B, 0x4D)
    # 배경색 설정
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F2F2F2')
    pPr.append(shd)

def parse_inline(text):
    """**bold** 와 `code` 파싱 → (text, bold, code) 세그먼트 리스트"""
    segments = []
    pattern = re.compile(r'(\*\*(.+?)\*\*|`([^`]+)`)')
    last = 0
    for m in pattern.finditer(text):
        if m.start() > last:
            segments.append((text[last:m.start()], False, False))
        if m.group(2):
            segments.append((m.group(2), True, False))
        elif m.group(3):
            segments.append((m.group(3), False, True))
        last = m.end()
    if last < len(text):
        segments.append((text[last:], False, False))
    return segments

def add_paragraph_with_inline(doc, text, style_name=None, indent=None):
    text = re.sub(r'^>\s*', '', text)  # blockquote 기호 제거
    p = doc.add_paragraph(style=style_name) if style_name else doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = indent
    for seg_text, bold, code in parse_inline(text):
        run = p.add_run(seg_text)
        run.font.name = '맑은 고딕'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        run.font.size = Pt(10)
        if bold:
            run.font.bold = True
        if code:
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
    return p

def add_section_break(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pgSzEl = OxmlElement('w:sectPr')
    pPr.append(pgSzEl)

def process_md_file(doc, filepath, is_first=False):
    with open(filepath, encoding='utf-8') as f:
        lines = f.readlines()

    if not is_first:
        doc.add_page_break()

    i = 0
    in_code = False
    code_lines = []

    while i < len(lines):
        line = lines[i].rstrip('\n')

        # 코드 블록 시작/끝
        if line.strip().startswith('```'):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                add_code_block(doc, code_lines)
                in_code = False
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        # 빈 줄
        if not line.strip():
            i += 1
            continue

        # 구분선
        if re.match(r'^---+$', line.strip()):
            p = doc.add_paragraph('─' * 60)
            p.runs[0].font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            p.runs[0].font.size = Pt(8)
            i += 1
            continue

        # 제목
        h_match = re.match(r'^(#{1,3})\s+(.*)', line)
        if h_match:
            level = len(h_match.group(1))
            title = h_match.group(2)
            p = doc.add_paragraph()
            run = p.add_run(title)
            set_heading_style(p, level)
            if level == 1:
                p.paragraph_format.space_before = Pt(18)
                p.paragraph_format.space_after = Pt(6)
            elif level == 2:
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(4)
            else:
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after = Pt(2)
            i += 1
            continue

        # 테이블
        if line.strip().startswith('|'):
            i = add_table_from_md(doc, lines, i)
            continue

        # 인용구
        if line.strip().startswith('>'):
            p = add_paragraph_with_inline(doc, line.strip(), indent=Cm(0.8))
            p.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55) if p.runs else None
            p.paragraph_format.space_after = Pt(2)
            i += 1
            continue

        # 리스트
        list_match = re.match(r'^(\s*)([-*]|\d+\.)\s+(.*)', line)
        if list_match:
            indent_level = len(list_match.group(1)) // 2
            content = list_match.group(3)
            bullet = '•' if re.match(r'[-*]', list_match.group(2)) else list_match.group(2)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5 + indent_level * 0.5)
            p.paragraph_format.space_after = Pt(1)
            p.add_run(f'{bullet} ')
            for seg_text, bold, code in parse_inline(content):
                run = p.add_run(seg_text)
                run.font.name = '맑은 고딕'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                run.font.size = Pt(10)
                if bold: run.font.bold = True
                if code:
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
            i += 1
            continue

        # 일반 텍스트
        add_paragraph_with_inline(doc, line.strip())
        i += 1

# ── 표지 ──────────────────────────────────────────────────────────
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_before = Pt(80)
run = title_p.add_run('Piuda')
run.font.size = Pt(36)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
run.font.name = '맑은 고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = sub_p.add_run('치매 케어 서비스 — 기획 문서')
run2.font.size = Pt(16)
run2.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
run2.font.name = '맑은 고딕'
run2._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

doc.add_paragraph()
date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = date_p.add_run('2026년 05월 30일')
run3.font.size = Pt(11)
run3.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
run3.font.name = '맑은 고딕'
run3._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

doc.add_page_break()

# ── 파일 처리 ──────────────────────────────────────────────────────
base = '/home/donghyun/바탕화면/Piuda/Server/BackEnd/piuda/docs'
files = [
    (f'{base}/기능명세서.md', True),
    (f'{base}/사용자시나리오.md', False),
    (f'{base}/API명세서.md', False),
]

for path, is_first in files:
    process_md_file(doc, path, is_first)

output = f'{base}/Piuda_기획문서.docx'
doc.save(output)
print(f'완료: {output}')
